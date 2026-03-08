#!/usr/bin/env python3
"""
NCBI Data Retriever
===================

A user-friendly tool for retrieving GenBank and FASTA files from NCBI databases
using accession IDs. This script supports multiple input file formats and is
designed for students and researchers who may not be familiar with programming.

Author: Bio357 Course
Version: 1.0
"""

import os
import sys
import time
import logging

# Try to import required modules with fallbacks
try:
    import yaml
except ImportError:
    print("Error: PyYAML is required. Install with: pip install PyYAML")
    sys.exit(1)

try:
    import requests
except ImportError:
    print("Error: requests is required. Install with: pip install requests")
    sys.exit(1)

try:
    from pathlib import Path
except ImportError:
    print("Error: pathlib is required (Python 3.4+). Please upgrade Python.")
    sys.exit(1)

# Optional typing support (Python 3.5+)
try:
    from typing import List, Dict, Any
except ImportError:
    # Fallback for older Python versions
    List = list
    Dict = dict
    Any = object

# File format support with error handling
try:
    import pandas as pd
except ImportError:
    print("Warning: pandas not available. CSV and Excel support disabled.")
    pd = None

try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not available. Word document support disabled.")
    Document = None

try:
    import openpyxl
except ImportError:
    print("Warning: openpyxl not available. Excel support may be limited.")
    openpyxl = None

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ncbi_retriever.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


class NCBIDataRetriever:
    """
    A class to retrieve GenBank and FASTA files from NCBI using accession IDs.
    
    This class handles multiple input file formats and provides comprehensive
    error handling and progress reporting.
    """
    
    def __init__(self, config_file: str = "config.yaml"):
        """
        Initialize the NCBI Data Retriever.
        
        Args:
            config_file (str): Path to the YAML configuration file
        """
        self.config = self._load_config(config_file)
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': f'NCBI_Data_Retriever/1.0 ({self.config["email"]})'
        })
        
        # Create output directory if it doesn't exist
        if self.config.get('output_path'):
            output_path = self.config['output_path']
            try:
                os.makedirs(output_path, exist_ok=True)
                self.output_dir = Path(output_path).resolve()
                logger.info(f"Output directory: {self.output_dir}")
            except Exception as e:
                logger.warning(f"Could not create output directory '{output_path}': {e}")
                logger.info("Using current directory instead")
                self.output_dir = Path.cwd()
        else:
            self.output_dir = Path.cwd()
            
        logger.info(f"NCBI Data Retriever initialized")
        logger.info(f"Working directory: {Path.cwd()}")
        logger.info(f"Output directory: {self.output_dir}")
    
    def _load_config(self, config_file: str) -> Dict[str, Any]:
        """
        Load configuration from YAML file.
        
        Args:
            config_file (str): Path to configuration file
            
        Returns:
            Dict[str, Any]: Configuration dictionary
        """
        try:
            with open(config_file, 'r') as file:
                config = yaml.safe_load(file)
            
            # Validate required fields
            required_fields = ['email', 'input_file_path']
            for field in required_fields:
                if field not in config:
                    raise ValueError(f"Required field '{field}' not found in config file")
            
            # Set defaults for optional fields
            config.setdefault('batch_size', 200)
            config.setdefault('delay_between_requests', 0.5)
            config.setdefault('download_genbank', True)
            config.setdefault('download_fasta', True)
            
            return config
            
        except FileNotFoundError:
            logger.error(f"Configuration file '{config_file}' not found!")
            sys.exit(1)
        except yaml.YAMLError as e:
            logger.error(f"Error parsing YAML file: {e}")
            sys.exit(1)
        except Exception as e:
            logger.error(f"Error loading configuration: {e}")
            sys.exit(1)
    
    def _read_accession_ids(self, file_path: str) -> List[str]:
        """
        Read accession IDs from various file formats.
        
        Supported formats:
        - .txt: Plain text files (one ID per line)
        - .csv: Comma-separated values
        - .xlsx: Excel files
        - .docx: Word documents
        
        Args:
            file_path (str): Path to the input file
            
        Returns:
            List[str]: List of accession IDs
        """
        file_path = Path(file_path)
        
        # Check if file exists and provide helpful error message
        if not file_path.exists():
            # Try to provide helpful suggestions
            current_dir = Path.cwd()
            logger.error(f"Input file '{file_path}' not found!")
            logger.error(f"Current working directory: {current_dir}")
            logger.error("Troubleshooting tips:")
            logger.error("1. Make sure you're using the FULL PATH to your file")
            logger.error("2. Check if the file exists in the specified location")
            logger.error("3. Try using an absolute path like:")
            logger.error(f"   - Mac/Linux: /Users/yourname/Documents/yourfile.txt")
            logger.error(f"   - Windows: C:\\Users\\yourname\\Documents\\yourfile.txt")
            logger.error("4. Or place your file in the ncbi_tools folder and use just the filename")
            raise FileNotFoundError(f"Input file '{file_path}' not found! See log for troubleshooting tips.")
        
        logger.info(f"Reading accession IDs from: {file_path}")
        logger.info(f"File size: {file_path.stat().st_size} bytes")
        
        try:
            if file_path.suffix.lower() == '.txt':
                return self._read_txt_file(file_path)
            elif file_path.suffix.lower() == '.csv':
                return self._read_csv_file(file_path)
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                return self._read_excel_file(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._read_docx_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error reading file '{file_path}': {e}")
            raise
    
    def _read_txt_file(self, file_path: Path) -> List[str]:
        """Read accession IDs from a plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        
        # Clean and filter lines
        accession_ids = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith('#'):  # Skip empty lines and comments
                accession_ids.append(line)
        
        return accession_ids
    
    def _read_csv_file(self, file_path: Path) -> List[str]:
        """Read accession IDs from a CSV file."""
        if pd is None:
            raise ImportError("pandas is required for CSV file support. Install with: pip install pandas")
        
        df = pd.read_csv(file_path)
        
        # Try to find a column with accession IDs
        possible_columns = ['accession', 'accession_id', 'id', 'acc', 'sequence_id']
        accession_column = None
        
        for col in possible_columns:
            if col.lower() in [c.lower() for c in df.columns]:
                accession_column = col
                break
        
        if accession_column is None:
            # If no obvious column found, use the first column
            accession_column = df.columns[0]
            logger.warning(f"No obvious accession column found. Using first column: {accession_column}")
        
        accession_ids = df[accession_column].dropna().astype(str).tolist()
        return [acc.strip() for acc in accession_ids if acc.strip()]
    
    def _read_excel_file(self, file_path: Path) -> List[str]:
        """Read accession IDs from an Excel file."""
        if pd is None:
            raise ImportError("pandas is required for Excel file support. Install with: pip install pandas")
        
        df = pd.read_excel(file_path)
        
        # Try to find a column with accession IDs
        possible_columns = ['accession', 'accession_id', 'id', 'acc', 'sequence_id']
        accession_column = None
        
        for col in possible_columns:
            if col.lower() in [c.lower() for c in df.columns]:
                accession_column = col
                break
        
        if accession_column is None:
            # If no obvious column found, use the first column
            accession_column = df.columns[0]
            logger.warning(f"No obvious accession column found. Using first column: {accession_column}")
        
        accession_ids = df[accession_column].dropna().astype(str).tolist()
        return [acc.strip() for acc in accession_ids if acc.strip()]
    
    def _read_docx_file(self, file_path: Path) -> List[str]:
        """Read accession IDs from a Word document."""
        if Document is None:
            raise ImportError("python-docx is required for Word document support. Install with: pip install python-docx")
        
        doc = Document(file_path)
        accession_ids = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and not text.startswith('#'):
                # Split by common delimiters
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        accession_ids.append(line)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and not text.startswith('#'):
                        accession_ids.append(text)
        
        return accession_ids
    
    def _clean_accession_ids(self, accession_ids: List[str]) -> List[str]:
        """
        Clean and validate accession IDs.
        
        Args:
            accession_ids (List[str]): Raw accession IDs
            
        Returns:
            List[str]: Cleaned and validated accession IDs
        """
        cleaned_ids = []
        
        for acc_id in accession_ids:
            # Remove common prefixes and clean up
            acc_id = acc_id.strip()
            acc_id = acc_id.replace('Accession:', '').replace('ACC:', '').replace('ID:', '')
            acc_id = acc_id.strip()
            
            # Basic validation (NCBI accession IDs are typically alphanumeric)
            if acc_id and acc_id.replace('_', '').replace('.', '').isalnum():
                cleaned_ids.append(acc_id)
            else:
                logger.warning(f"Skipping invalid accession ID: {acc_id}")
        
        # Remove duplicates while preserving order
        seen = set()
        unique_ids = []
        for acc_id in cleaned_ids:
            if acc_id not in seen:
                seen.add(acc_id)
                unique_ids.append(acc_id)
        
        logger.info(f"Found {len(unique_ids)} unique accession IDs")
        return unique_ids
    
    def _download_from_ncbi(self, accession_ids: List[str], file_format: str) -> Dict[str, str]:
        """
        Download sequences from NCBI using Entrez API.
        
        Args:
            accession_ids (List[str]): List of accession IDs
            file_format (str): Format to download ('genbank' or 'fasta')
            
        Returns:
            Dict[str, str]: Dictionary mapping accession IDs to file paths
        """
        base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
        downloaded_files = {}
        
        # Process in batches
        batch_size = self.config.get('batch_size', 200)
        delay = self.config.get('delay_between_requests', 0.5)
        
        for i in range(0, len(accession_ids), batch_size):
            batch = accession_ids[i:i + batch_size]
            logger.info(f"Processing batch {i//batch_size + 1}/{(len(accession_ids)-1)//batch_size + 1} for {file_format} format")
            logger.info(f"Batch contains {len(batch)} accession IDs: {batch[:3]}{'...' if len(batch) > 3 else ''}")
            
            try:
                # Step 1: Search for the accession IDs
                search_url = f"{base_url}esearch.fcgi"
                search_params = {
                    'db': 'nucleotide',
                    'term': ' OR '.join(batch),
                    'retmode': 'json',
                    'retmax': len(batch)
                }
                
                response = self.session.get(search_url, params=search_params)
                response.raise_for_status()
                search_data = response.json()
                
                if 'esearchresult' not in search_data or not search_data['esearchresult']['idlist']:
                    logger.warning(f"No results found for batch starting at index {i}")
                    continue
                
                # Step 2: Fetch the sequences
                fetch_url = f"{base_url}efetch.fcgi"
                
                # Map file format to correct NCBI rettype parameter
                if file_format == 'genbank':
                    rettype = 'gb'
                elif file_format == 'fasta':
                    rettype = 'fasta'
                else:
                    rettype = file_format
                
                fetch_params = {
                    'db': 'nucleotide',
                    'id': ','.join(search_data['esearchresult']['idlist']),
                    'rettype': rettype,
                    'retmode': 'text'
                }
                
                logger.info(f"Fetching {file_format} format using rettype={rettype}")
                
                response = self.session.get(fetch_url, params=fetch_params)
                response.raise_for_status()
                
                # Check if we got actual data
                response_text = response.text.strip()
                if not response_text:
                    logger.warning(f"No data received for {file_format} format in batch {i//batch_size + 1}")
                    continue
                
                # Check for error messages in response
                if "Error" in response_text or "error" in response_text:
                    logger.warning(f"Error in {file_format} response: {response_text[:200]}...")
                    continue
                
                # Save the batch file with appropriate extension
                if file_format == 'genbank':
                    file_extension = 'genbank'
                elif file_format == 'fasta':
                    file_extension = 'fasta'
                else:
                    file_extension = file_format
                
                batch_filename = f"batch_{i//batch_size + 1}_{file_format}.{file_extension}"
                batch_filepath = self.output_dir / batch_filename
                
                with open(batch_filepath, 'w', encoding='utf-8') as f:
                    f.write(response_text)
                
                logger.info(f"✅ Downloaded {len(search_data['esearchresult']['idlist'])} sequences to {batch_filepath}")
                logger.info(f"   File size: {len(response_text)} characters")
                
                # Map accession IDs to file paths
                for acc_id in batch:
                    downloaded_files[acc_id] = str(batch_filepath)
                
                # Be respectful to NCBI servers
                time.sleep(delay)
                
            except requests.exceptions.RequestException as e:
                logger.error(f"Error downloading batch {i//batch_size + 1}: {e}")
                continue
            except Exception as e:
                logger.error(f"Unexpected error processing batch {i//batch_size + 1}: {e}")
                continue
        
        return downloaded_files
    
    def retrieve_data(self) -> None:
        """
        Main method to retrieve GenBank and FASTA files from NCBI.
        """
        try:
            # Read accession IDs from input file
            raw_accession_ids = self._read_accession_ids(self.config['input_file_path'])
            accession_ids = self._clean_accession_ids(raw_accession_ids)
            
            if not accession_ids:
                logger.error("No valid accession IDs found!")
                return
            
            logger.info(f"Starting download for {len(accession_ids)} accession IDs")
            
            # Download GenBank files if requested
            if self.config.get('download_genbank', True):
                logger.info("=" * 50)
                logger.info("DOWNLOADING GENBANK FILES")
                logger.info("=" * 50)
                genbank_files = self._download_from_ncbi(accession_ids, 'genbank')
                logger.info(f"✅ Downloaded {len(genbank_files)} GenBank files")
                if genbank_files:
                    logger.info(f"GenBank files saved to: {list(genbank_files.values())[0]}")
            else:
                logger.info("⏭️  Skipping GenBank download (disabled in config)")
            
            # Download FASTA files if requested
            if self.config.get('download_fasta', True):
                logger.info("=" * 50)
                logger.info("DOWNLOADING FASTA FILES")
                logger.info("=" * 50)
                fasta_files = self._download_from_ncbi(accession_ids, 'fasta')
                logger.info(f"✅ Downloaded {len(fasta_files)} FASTA files")
                if fasta_files:
                    logger.info(f"FASTA files saved to: {list(fasta_files.values())[0]}")
            else:
                logger.info("⏭️  Skipping FASTA download (disabled in config)")
            
            logger.info("Download completed successfully!")
            
        except Exception as e:
            logger.error(f"Error during data retrieval: {e}")
            raise


def main():
    """
    Main function to run the NCBI Data Retriever.
    """
    print("=" * 60)
    print("NCBI Data Retriever")
    print("==================")
    print("A tool for downloading GenBank and FASTA files from NCBI")
    print("=" * 60)
    
    try:
        # Initialize the retriever
        retriever = NCBIDataRetriever()
        
        # Start the download process
        retriever.retrieve_data()
        
        print("\n" + "=" * 60)
        print("Download completed successfully!")
        print("Check the output directory for your downloaded files.")
        print("=" * 60)
        
    except KeyboardInterrupt:
        print("\nDownload interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nError: {e}")
        print("Please check the log file 'ncbi_retriever.log' for more details.")
        sys.exit(1)


if __name__ == "__main__":
    main()